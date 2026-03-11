import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

class Validator:
    def __init__(self, filepath: str):
        self.filepath = filepath
        self.doc = Document(filepath)
        self.errors = []
        self.errors_eng = []

    def validate(self):
        self.check_margins()
        self.check_orientation()
        self.check_font_and_size()
        self.check_line_spacing()
        self.check_first_line_indent()
        self.check_email()
        self.check_title()
        self.check_keywords()
        self.check_annotation()
        self.check_references_count()
        self.check_reference_format()
        return self.errors, self.errors_eng

    # Поля
    def check_margins(self):
        for section in self.doc.sections:
            if round(section.top_margin.cm, 1) != 2.5:
                self.errors.append("Отступ сверху должен равняться 2.5 единицам")
                self.errors_eng.append("The top margin must be 2.5 units")
            if round(section.bottom_margin.cm, 1) != 2.5:
                self.errors.append("Отступ снизу должен равняться 2.5 единицам")
                self.errors_eng.append("The bottom margin must be 2.5 units")
            if round(section.left_margin.cm, 1) != 2.5:
                self.errors.append("Отступ слева должен равняться 2.5 единицам")
                self.errors_eng.append("The left margin must be 2.5 units")
            if round(section.right_margin.cm, 1) != 2.5:
                self.errors.append("Отступ справа должен равняться 2.5 единицам")
                self.errors_eng.append("The right margin must be 2.5 units")

    # Ориентация страниц
    def check_orientation(self):
        for section in self.doc.sections:
            if section.orientation != 0:  # 0 = portrait
                self.errors.append("Ориентация страниц должна быть книжной")
                self.errors_eng.append("The orientation of the pages must be portrait")

    # Шрифт и размер
    def check_font_and_size(self):
        for p in self.doc.paragraphs:
            for run in p.runs:
                if run.font.name and run.font.name != "Times New Roman":
                    self.errors.append("В тексте статьи необходимо использовать шрифт Times New Roman, за исключением математических формул")
                    self.errors_eng.append("The Times New Roman font should be used in the text of the article, with the exception of mathematical formulas")
                    return
                if run.font.size and run.font.size.pt != 14:
                    self.errors.append("Размер шрифта должен быть 14")
                    self.errors_eng.append("The font size must be 14")
                    return

    # Межстрочный интервал
    def check_line_spacing(self):
        for p in self.doc.paragraphs:
            if p.paragraph_format.line_spacing:
                if p.paragraph_format.line_spacing != 1.5:
                    self.errors.append("Межстрочный интервал должен равняться 1.5 единицам")
                    self.errors_eng.append("The line spacing must be 1.5 units")
                    return

    # Абзацный отступ
    def check_first_line_indent(self):
        for p in self.doc.paragraphs:
            indent = p.paragraph_format.first_line_indent
            if indent and round(indent.cm, 2) != 1.25:
                self.errors.append("Абзацный отступ должен равняться 1.25 единицам")
                self.errors_eng.append("The paragraph indentation must be 1.25 units")
                return

    # Email (курсив)
    def check_email(self):
        email_pattern = r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}"
        found = False

        for p in self.doc.paragraphs:
            if re.search(email_pattern, p.text):
                found = True
                for run in p.runs:
                    if re.search(email_pattern, run.text):
                        if not run.italic:
                            self.errors.append("Email должен быть напечатан курсивом")
                            self.errors_eng.append("The email must be printed in italics")
                break

        if not found:
            self.errors.append("Email автора не найден")
            self.errors_eng.append("The author's email was not found")

    # Название статьи (по центру, жирное)
    def check_title(self):
        for p in self.doc.paragraphs[:10]:
            if p.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                bold_found = any(run.bold for run in p.runs)
                if not bold_found:
                    self.errors.append("Для названия статьи необходимо использовать полужирное начертание")
                    self.errors_eng.append("The title of the article must be printed in bold")
                return
        self.errors.append("Название статьи не найдено")
        self.errors_eng.append("The title of the article was not found")

    # Ключевые слова
    def check_keywords(self):
        for p in self.doc.paragraphs:
            if "key words" in p.text.lower():
                parts = p.text.split(":")
                if len(parts) > 1:
                    keywords = [k.strip() for k in parts[1].split(";")]
                    if not (5 <= len(keywords) <= 10):
                        self.errors.append("Ключевых слов должно быть от 5 до 10")
                        self.errors_eng.append("There should be from 5 to 10 keywords")
                return
        self.errors.append("Ключевые слова не найдены")
        self.errors_eng.append("Keywords of the article were not found")

    # Аннотация
    def check_annotation(self):
        for p in self.doc.paragraphs:
            if "abstract." in p.text.lower():
                abstract_text = p.text.lower().strip()
                abstract_text = abstract_text.replace(" ", "")
                text_len = len(abstract_text) - len("abstract.")
                print(text_len)
                if not (300 <= text_len <= 500):
                    self.errors.append("Аннотация на английском должна содержать от 300 до 500 знаков без учёта пробелов")
                    self.errors_eng.append("The abstract in English should contain from 300 to 500 characters without spaces")
                italic = any(run.italic for run in p.runs)
                if not italic:
                    self.errors.append("Аннотация должна быть напечатана курсивом")
                    self.errors_eng.append("The abstract should be printed in italics")
                return
        self.errors.append("Аннотация не найдена")
        self.errors_eng.append("The abstract was not found")

    # Формат ссылок
    def check_reference_format(self):
        pattern = r"\[\d+(,\s?с\.\s?\d+)?(;\s?\d+(,\s?с\.\s?\d+)?)?\]"
        for p in self.doc.paragraphs:
            if p.text == "REFERENCES":
                return
            matches = re.findall(r"\[.*?]", p.text)
            for m in matches:
                if not re.match(pattern, m):
                    self.errors.append(f"Неверный формат ссылки: {m}")
                    self.errors_eng.append(f"Invalid reference format: {m}")

    # Список литературы
    def check_references_count(self):
        start = False
        count = 0
        for p in self.doc.paragraphs:
            lower_text = p.text.lower()
            if "references" in lower_text or "список литературы" in lower_text: # TODO: нужно ли на английском?
                start = True
                continue
            if start:
                if p.text.strip() == "":
                    break
                count += 1

        if count < 5:
            self.errors.append("Список литературы должен содержать не менее 5 источников")
            self.errors_eng.append("The list of references should contain at least 5 sources.")
