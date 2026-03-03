import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

class Validator:
    def __init__(self, filepath: str):
        self.filepath = filepath
        self.doc = Document(filepath)
        self.errors = []

    def validate(self):
        self.check_margins()
        self.check_orientation()
        self.check_font_and_size()
        self.check_line_spacing()
        self.check_first_line_indent()
        self.check_email()
        self.check_title()
        self.check_keywords()
        self.check_annotation()
        self.check_references_count()
        self.check_reference_format()
        return self.errors

    # Поля
    def check_margins(self):
        for section in self.doc.sections:
            if round(section.top_margin.cm, 1) != 2.5:
                self.errors.append("Неправильный отступ сверху")
            if round(section.bottom_margin.cm, 1) != 2.5:
                self.errors.append("Неправильный отступ снизу")
            if round(section.left_margin.cm, 1) != 2.5:
                self.errors.append("Неправильный отступ слева")
            if round(section.right_margin.cm, 1) != 2.5:
                self.errors.append("Неправильный отступ справа")

    # Ориентация страниц
    def check_orientation(self):
        for section in self.doc.sections:
            if section.orientation != 0:  # 0 = portrait
                self.errors.append("Ориентация страниц должна быть книжной")

    # Шрифт и размер
    def check_font_and_size(self):
        for p in self.doc.paragraphs:
            for run in p.runs:
                if run.font.name and run.font.name != "Times New Roman":
                    self.errors.append("Используется не Times New Roman")
                    return
                if run.font.size and run.font.size.pt != 14:
                    self.errors.append("Размер шрифта должен быть 14")
                    return

    # Межстрочный интервал
    def check_line_spacing(self):
        for p in self.doc.paragraphs:
            if p.paragraph_format.line_spacing:
                if p.paragraph_format.line_spacing != 1.5:
                    self.errors.append("Межстрочный интервал должен быть 1.5")
                    return

    # Абзацный отступ
    def check_first_line_indent(self):
        for p in self.doc.paragraphs:
            indent = p.paragraph_format.first_line_indent
            if indent and round(indent.cm, 2) != 1.25:
                self.errors.append("Абзацный отступ должен быть 1.25")
                return

    # Email (курсив)
    def check_email(self):
        email_pattern = r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}"
        found = False

        for p in self.doc.paragraphs:
            if re.search(email_pattern, p.text):
                found = True
                for run in p.runs:
                    if re.search(email_pattern, run.text):
                        if not run.italic:
                            self.errors.append("Email должен быть курсивом")
                break

        if not found:
            self.errors.append("Email не найден")

    # Название статьи (по центру, жирное)
    def check_title(self):
        for p in self.doc.paragraphs[:10]:
            if p.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                bold_found = any(run.bold for run in p.runs)
                if not bold_found:
                    self.errors.append("Название статьи должно быть жирным")
                return
        self.errors.append("Название статьи не найдено")

    # Ключевые слова
    def check_keywords(self):
        for p in self.doc.paragraphs:
            if "key words" in p.text.lower():
                parts = p.text.split(":")
                if len(parts) > 1:
                    keywords = [k.strip() for k in parts[1].split(";")]
                    if not (5 <= len(keywords) <= 10):
                        self.errors.append("Ключевых слов должно быть 5–10")
                return
        self.errors.append("Ключевые слова не найдены")

    # Аннотация
    def check_annotation(self):
        for p in self.doc.paragraphs:
            if "abstract." in p.text.lower():
                text_len = len(p.text.lower().strip()) - len("abstract. ")
                print(text_len)
                if not (300 <= text_len <= 500):
                    self.errors.append("Аннотация должна быть 300–500 знаков")
                italic = any(run.italic for run in p.runs)
                if not italic:
                    self.errors.append("Аннотация должна быть напечатана курсивом")
                return
        self.errors.append("Аннотация не найдена")

    # Формат ссылок
    def check_reference_format(self):
        pattern = r"\[\d+(,\s?с\.\s?\d+)?(;\s?\d+(,\s?с\.\s?\d+)?)?\]"
        for p in self.doc.paragraphs:
            matches = re.findall(r"\[.*?\]", p.text)
            for m in matches:
                if not re.match(pattern, m):
                    self.errors.append(f"Неверный формат ссылки: {m}")

    # Список литературы
    def check_references_count(self):
        start = False
        count = 0
        for p in self.doc.paragraphs:
            if "список литературы" in p.text.lower(): # TODO: нужно ли на английском?
                start = True
                continue
            if start:
                if p.text.strip() == "":
                    break
                count += 1

        if count < 5:
            self.errors.append("Список литературы должен содержать не менее 5 источников")
